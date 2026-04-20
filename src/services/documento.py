"""
Generación de documentos Word y conversión a PDF.
Adaptado de agent-old-dontlook/cotizacion-generator/scripts/generar_cotizacion.py.
"""
import io
import logging
import os
import subprocess
import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor
from sqlalchemy import select
from sqlalchemy.orm import selectinload

from src.db.engine import async_session_factory
from src.db.models import Cotizacion, LineaCotizacion
from src.services.cotizacion import formato_cop_sin_signo

logger = logging.getLogger(__name__)

TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "cotizacion_v2.docx"
CONTRATO_TEMPLATE_PATH = Path(__file__).parent.parent / "templates" / "contrato_v1.docx"

DATOS_BANCARIOS = {
    "nit_ml": "900.000.000-0",   # actualizar con datos reales
    "banco": "Bancolombia",
    "tipo_cuenta": "Cuenta Corriente",
    "numero_cuenta": "000-000000-00",  # actualizar con datos reales
}

MESES_ES = {
    1: "enero", 2: "febrero", 3: "marzo", 4: "abril",
    5: "mayo", 6: "junio", 7: "julio", 8: "agosto",
    9: "septiembre", 10: "octubre", 11: "noviembre", 12: "diciembre",
}


# ─── Funciones públicas ────────────────────────────────────────────────────────

async def generar_word(cotizacion_id: int) -> bytes:
    """
    Carga la cotización de la DB, llena la plantilla Word y retorna los bytes del .docx.
    """
    async with async_session_factory() as db:
        result = await db.execute(
            select(Cotizacion)
            .options(
                selectinload(Cotizacion.lineas).selectinload(LineaCotizacion.servicio),
                selectinload(Cotizacion.lineas).selectinload(LineaCotizacion.equipo_alquiler),
                selectinload(Cotizacion.cliente),
                selectinload(Cotizacion.contacto),
            )
            .where(Cotizacion.id == cotizacion_id)
        )
        cot = result.scalar_one()

    cliente = cot.cliente
    contacto = cot.contacto
    numero_limpio = cot.numero_cotizacion.replace("COT-", "")
    fecha_str = (
        f"{cot.fecha.day} de {MESES_ES[cot.fecha.month]} de {cot.fecha.year}"
    )
    nombre_completo = contacto.nombre_completo if contacto else (cliente.nombre_empresa if cliente else "")
    nombre_corto = nombre_completo.split()[0] if nombre_completo else ""

    # Determinar tipo de evento (presencial / virtual)
    tipo_evento = "virtual"
    for linea in cot.lineas:
        desc = (linea.descripcion_generada or "").lower()
        if "presencial" in desc:
            tipo_evento = "presencial"
            break

    doc = Document(str(TEMPLATE_PATH))

    # ── 1. Encabezado ──────────────────────────────────────────────────────────
    _replace(doc, "COT-___", f"COT-{numero_limpio}")
    _replace(doc, "DD de mes de AAAA", fecha_str)
    _replace(doc, "{nombre_contacto}", nombre_completo)
    _replace(doc, "{cargo}", contacto.cargo if contacto else "")
    _replace(doc, "{empresa}", cliente.nombre_empresa if cliente else "")
    _replace(doc, "{email}", contacto.email if contacto else "")
    _replace(doc, "{telefono}", contacto.telefono if contacto else "")
    _replace(doc, "{nombre_corto}", nombre_corto)
    _replace(doc, "{referencia_servicio}", _referencia_servicio(cot.lineas))
    _replace(doc, "{presencial/virtual}", tipo_evento)
    _replace(doc, "{ubicacion_evento}", cot.ubicacion_evento or "Bogotá")

    # ── 2. Tabla de servicios ──────────────────────────────────────────────────
    if doc.tables:
        services_table = doc.tables[0]
        lineas_ord = sorted(cot.lineas, key=lambda l: l.orden)
        for i, linea in enumerate(lineas_ord[:5]):
            row = services_table.rows[i + 2]
            servicio = linea.servicio
            _celda(row.cells[0], str(i + 1), size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _celda(row.cells[1], linea.descripcion_generada or (servicio.nombre if servicio else ""), size=7, bold=True)
            _celda(row.cells[2], servicio.idioma_origen or "" if servicio else "", size=7)
            _celda(row.cells[3], servicio.idioma_destino or "" if servicio else "", size=7)
            _celda(row.cells[4], _formato_fecha_rango(linea.fecha_servicio_inicio, linea.fecha_servicio_fin), size=7)
            _celda(row.cells[5], linea.horario or "", size=7)
            _celda(row.cells[6], str(int(linea.cantidad)), size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _celda(row.cells[7], _unidad_label(servicio), size=7)
            _celda(row.cells[8], formato_cop_sin_signo(linea.precio_unitario), size=7, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            _celda(row.cells[9], formato_cop_sin_signo(linea.precio_total), size=7, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
            _celda(row.cells[10], "", size=6)

    # ── 3. Totales ─────────────────────────────────────────────────────────────
    _replace(doc, "{subtotal}", formato_cop_sin_signo(cot.subtotal))
    _replace(doc, "{iva}", "EXENTO" if cot.exento_iva else formato_cop_sin_signo(cot.iva))
    _replace(doc, "{total}", formato_cop_sin_signo(cot.total))

    # ── 4. Condiciones ─────────────────────────────────────────────────────────
    condiciones = _generar_condiciones(cot)
    for i in range(3):
        _replace(doc, f"{{notas_servicio_{i+1}}}", condiciones["notas"][i] if i < len(condiciones["notas"]) else "")
    _replace(doc, "{validez_oferta}", condiciones["validez_oferta"])
    _replace(doc, "{forma_pago}", condiciones["forma_pago"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


async def generar_word_contrato(cotizacion_id: int) -> bytes:
    """
    Genera el documento Word de confirmación de servicio (contrato) a partir
    de una cotización aprobada. Retorna los bytes del .docx.
    """
    import datetime
    async with async_session_factory() as db:
        result = await db.execute(
            select(Cotizacion)
            .options(
                selectinload(Cotizacion.lineas).selectinload(LineaCotizacion.servicio),
                selectinload(Cotizacion.cliente),
                selectinload(Cotizacion.contacto),
            )
            .where(Cotizacion.id == cotizacion_id)
        )
        cot = result.scalar_one()

    cliente = cot.cliente
    contacto = cot.contacto
    numero_cot = cot.numero_cotizacion
    numero_cont = numero_cot.replace("COT-", "CONT-")
    fecha_str = (
        f"{cot.fecha.day} de {MESES_ES[cot.fecha.month]} de {cot.fecha.year}"
        if cot.fecha else datetime.date.today().strftime("%d/%m/%Y")
    )

    linea_principal = sorted(cot.lineas, key=lambda l: l.orden)[0] if cot.lineas else None
    servicio = linea_principal.servicio if linea_principal else None

    descripcion_servicio = linea_principal.descripcion_generada if linea_principal else "Servicio lingüístico"
    fecha_inicio_str = _formato_fecha_rango(
        linea_principal.fecha_servicio_inicio if linea_principal else None,
        linea_principal.fecha_servicio_fin if linea_principal else None,
    ) if linea_principal else ""
    fecha_fin_str = ""
    horario_str = linea_principal.horario or "" if linea_principal else ""
    num_interpretes_str = (
        f"{linea_principal.num_interpretes} intérprete(s)" if linea_principal and linea_principal.num_interpretes
        else "Según servicio"
    )

    anticipo = (cot.total or 0) / 2
    saldo = (cot.total or 0) / 2

    doc = Document(str(CONTRATO_TEMPLATE_PATH))

    _replace(doc, "{numero_contrato}", numero_cont)
    _replace(doc, "{fecha}", fecha_str)
    _replace(doc, "{numero_cotizacion}", numero_cot)
    _replace(doc, "{empresa}", cliente.nombre_empresa if cliente else "")
    _replace(doc, "{nit}", cliente.nit or "Por confirmar" if cliente else "Por confirmar")
    _replace(doc, "{nombre_contacto}", contacto.nombre_completo if contacto else "")
    _replace(doc, "{cargo}", contacto.cargo or "" if contacto else "")
    _replace(doc, "{email}", contacto.email or "" if contacto else "")
    _replace(doc, "{telefono}", contacto.telefono or "" if contacto else "")
    _replace(doc, "{descripcion_servicio}", descripcion_servicio)
    _replace(doc, "{fecha_inicio}", fecha_inicio_str)
    _replace(doc, "{fecha_fin}", fecha_fin_str)
    _replace(doc, "{fecha_inicio} — {fecha_fin}", fecha_inicio_str)
    _replace(doc, "{horario}", horario_str)
    _replace(doc, "{ubicacion}", cot.ubicacion_evento or "Bogotá")
    _replace(doc, "{num_interpretes}", num_interpretes_str)
    _replace(doc, "{subtotal}", formato_cop_sin_signo(cot.subtotal or 0))
    _replace(doc, "{iva}", "EXENTO" if cot.exento_iva else formato_cop_sin_signo(cot.iva or 0))
    _replace(doc, "{total}", formato_cop_sin_signo(cot.total or 0))
    _replace(doc, "{anticipo}", formato_cop_sin_signo(anticipo))
    _replace(doc, "{saldo}", formato_cop_sin_signo(saldo))
    _replace(doc, "{validez_oferta}", cot.validez_oferta or "30 días calendario")
    _replace(doc, "{nit_ml}", DATOS_BANCARIOS["nit_ml"])
    _replace(doc, "{banco}", DATOS_BANCARIOS["banco"])
    _replace(doc, "{tipo_cuenta}", DATOS_BANCARIOS["tipo_cuenta"])
    _replace(doc, "{numero_cuenta}", DATOS_BANCARIOS["numero_cuenta"])

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_a_pdf(docx_bytes: bytes) -> bytes:
    """
    Convierte bytes de un .docx a bytes de un .pdf usando LibreOffice headless.
    LibreOffice está instalado en el Dockerfile (libreoffice-writer).
    """
    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "cotizacion.docx")
        pdf_path = os.path.join(tmp, "cotizacion.pdf")

        with open(docx_path, "wb") as f:
            f.write(docx_bytes)

        # Buscar ejecutable de LibreOffice
        import shutil
        lo_cmd = shutil.which("libreoffice") or shutil.which("soffice")
        if not lo_cmd:
            raise RuntimeError("LibreOffice no encontrado. Verificar instalación en Docker.")

        subprocess.run(
            [lo_cmd, "--headless", "--convert-to", "pdf", "--outdir", tmp, docx_path],
            check=True,
            timeout=45,
            capture_output=True,
        )

        if not os.path.exists(pdf_path):
            raise RuntimeError("LibreOffice no generó el PDF. Verificar logs del proceso.")

        with open(pdf_path, "rb") as f:
            return f.read()


# ─── Helpers privados ──────────────────────────────────────────────────────────

def _replace(doc: Document, placeholder: str, value: str) -> None:
    """Reemplaza placeholder en todos los párrafos y celdas del documento."""
    value = str(value) if value else ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)


def _celda(cell, texto: str, size: int = 8, bold: bool = False,
           color_hex: str = "404040", alignment=None) -> None:
    """Escribe texto en una celda preservando el formato de ML Traductores."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(str(texto))
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16),
    )
    run.font.bold = bold
    run.font.name = "Calibri"


def _unidad_label(servicio) -> str:
    if servicio is None:
        return ""
    labels = {
        "por_hora": "Hora(s)",
        "por_dia": "Día(s)",
        "por_palabra": "Palabra(s)",
        "por_minuto": "Minuto(s)",
        "por_pagina": "Página(s)",
        "por_evento": "Evento",
    }
    return labels.get(servicio.unidad_cobro, servicio.unidad_cobro)


def _formato_fecha(d) -> str:
    if d is None:
        return ""
    return f"{d.day} {MESES_ES.get(d.month, '')} {d.year}"


def _formato_fecha_rango(inicio, fin) -> str:
    """Formatea rango de fechas: '28-29 abr', '28 abr - 3 may', '28 abr 2026'."""
    if inicio is None:
        return ""
    mes_i = MESES_ES.get(inicio.month, "")[:3]
    if fin is None or fin == inicio:
        return f"{inicio.day} {mes_i}"
    mes_f = MESES_ES.get(fin.month, "")[:3]
    if inicio.month == fin.month:
        return f"{inicio.day}-{fin.day} {mes_i}"
    return f"{inicio.day} {mes_i} - {fin.day} {mes_f}"


def _referencia_servicio(lineas) -> str:
    """Genera una descripción breve del evento para el encabezado."""
    if not lineas:
        return "Servicios lingüísticos"
    l = sorted(lineas, key=lambda x: x.orden)[0]
    return l.descripcion_generada or "Servicios lingüísticos"


def _detectar_tipo_servicio(cot: Cotizacion) -> dict:
    """Detecta el tipo de servicio desde las líneas de cotización."""
    es_presencial = False
    es_documentos = False
    for linea in cot.lineas:
        svc = linea.servicio
        if not svc:
            continue
        nombre = (svc.nombre or "").lower()
        cat = (svc.categoria or "").lower()
        if svc.es_presencial:
            es_presencial = True
        if (
            any(kw in nombre or kw in cat for kw in ["document", "traducc", "escrit"])
            or svc.unidad_cobro in ("por_palabra", "por_pagina")
        ):
            es_documentos = True
    return {
        "es_presencial": es_presencial,
        "es_documentos": es_documentos,
        "es_fuera_de_bogota": bool(cot.es_fuera_de_bogota),
        "incluir_terminos_corporativos": getattr(cot, "incluir_terminos_corporativos", False),
    }


_FORMA_PAGO_CORPORATIVA = (
    "Anticipo del 50% antes del inicio del evento.\n"
    "Saldo 50% Contado Comercial al finalizar el evento contra presentación de informe y factura.\n\n"
    "Términos y condiciones:\n"
    "- INCLUYE COORDINADOR ASIGNADO.\n"
    "- La presente oferta no incluye ningún suministro no especificado dentro de esta.\n"
    "- Toda cancelación una vez aceptada la oferta conlleva al pago del 50%.\n"
    "- No se asigna cotización parcial, incluye: Intérpretes simultáneos, "
    "Intérpretes de señas, equipos de traducción.\n"
    "- Cronograma: Montaje y pruebas: Desde 24 horas antes del inicio del evento.\n"
    "- Pólizas: A solicitud del cliente."
)

_VALIDEZ_INTERPRETACION = (
    "CONFIRMACIÓN ___ días antes del evento con el fin de reservar "
    "la agenda de los intérpretes asignados."
)


def _generar_condiciones(cot: Cotizacion) -> dict:
    """
    Genera notas, validez y forma de pago según el tipo de servicio.
    Retorna: {"notas": [s1, s2, s3], "validez_oferta": str, "forma_pago": str}
    """
    tipo = _detectar_tipo_servicio(cot)

    forma_pago_base = cot.forma_pago or "50% anticipo, 50% al finalizar"
    forma_pago = _FORMA_PAGO_CORPORATIVA if tipo["incluir_terminos_corporativos"] else forma_pago_base

    if tipo["es_documentos"]:
        notas = [
            (
                "Proceso de entrega:\n"
                "- Primera entrega para consideración del cliente, se reciben los comentarios con control de cambios.\n"
                "- Si hay comentarios se hacen ajustes (de forma, unificación de vocabulario, etc)\n"
                "- Segunda y última entrega del documento ajustado.\n"
                "- NUESTRO EQUIPO DE TRABAJO ESTÁ CONFORMADO POR TRADUCTORES OFICIALES DEBIDAMENTE ACREDITADOS.\n"
                "- Entrega del documento en Word o mismo formato si el archivo recibido es editable."
            ),
            "",
            "",
        ]
        validez = cot.validez_oferta or "30 días calendario"

    elif tipo["es_presencial"]:
        nota_equipos = (
            "- Un día de equipos se entiende de ocho horas o cualquier fracción del mismo\n"
            "- Transporte/Instalación/desmantelamiento (Incluido)\n"
            "- Montaje mismo día 6am"
        )
        nota_fuera = (
            "- Abono del 25%\n"
            "- Transporte: Viáticos: ___________"
        ) if tipo["es_fuera_de_bogota"] else ""

        notas = [
            (
                "Para intervenciones tipo conferencia se requiere la prestación del servicio de "
                "interpretación por un equipo de dos intérpretes en cabina a partir de 1.5 horas de servicio."
            ),
            nota_equipos,
            nota_fuera,
        ]
        validez = cot.validez_oferta or _VALIDEZ_INTERPRETACION

    else:  # virtual / remota (default)
        notas = [
            "Conexión 30 minutos antes del inicio del evento",
            "En caso de requerir conexión antes o pruebas anteriores favor adicionar $400.000 más IVA",
            "",
        ]
        validez = cot.validez_oferta or _VALIDEZ_INTERPRETACION

    return {"notas": notas, "validez_oferta": validez, "forma_pago": forma_pago}
