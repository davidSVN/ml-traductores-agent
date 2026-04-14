#!/usr/bin/env python3
"""
Generador de Cotizaciones ML Traductores
=========================================
Recibe un JSON con datos ya calculados y aprobados,
llena el template Word y convierte a PDF.

Este script NO consulta la DB ni toma decisiones de precio.
Eso lo hace el agente ANTES de llamar este script.

Uso:
    python generar_cotizacion.py --datos '{"numero_cotizacion":"050",...}'
    python generar_cotizacion.py --archivo datos_cotizacion.json

Requisitos:
    pip install python-docx --break-system-packages
"""

import json
import sys
import os
import argparse
import subprocess
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx no instalado.")
    print("  pip install python-docx --break-system-packages")
    sys.exit(1)


def formato_cop(monto):
    """$1.600.300 — formato colombiano sin decimales."""
    return f"${int(monto):,}".replace(",", ".")


def formato_cop_sin_signo(monto):
    """1.600.300 — sin signo $ (para celdas que ya tienen $)."""
    return f"{int(monto):,}".replace(",", ".")


def limpiar_texto_celda(cell, texto, size=8, bold=False, color_hex="404040", alignment=None):
    """Escribe texto en una celda de tabla preservando formato."""
    for p in cell.paragraphs:
        p.clear()
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(str(texto))
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(
        int(color_hex[0:2], 16),
        int(color_hex[2:4], 16),
        int(color_hex[4:6], 16)
    )
    run.font.bold = bold
    run.font.name = "Calibri"


def replace_placeholder(doc, placeholder, value):
    """Reemplaza placeholder en párrafos y celdas de tablas."""
    value = str(value) if value else ""
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, value)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, value)


def buscar_template(template_path):
    """Busca el template Word en varias ubicaciones posibles."""
    candidatos = [
        template_path,
        os.path.join(os.path.dirname(__file__), "..", "assets", "Formato_Cotizaciones_v2.docx"),
        "/mnt/user-data/outputs/Formato_Cotizaciones_v2.docx",
        "assets/Formato_Cotizaciones_v2.docx",
        "Formato_Cotizaciones_v2.docx",
    ]
    for ruta in candidatos:
        if ruta and os.path.exists(ruta):
            return ruta
    return None


def convertir_a_pdf(docx_path, output_dir):
    """Convierte Word a PDF. Intenta LibreOffice, luego docx2pdf."""
    pdf_path = docx_path.replace(".docx", ".pdf")

    # Método 1: LibreOffice
    try:
        import shutil
        _lo_candidates = [
            shutil.which("libreoffice"),
            shutil.which("soffice"),
            r"C:\Program Files\LibreOffice\program\soffice.exe",
            r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
        ]
        cmd = next((c for c in _lo_candidates if c and os.path.exists(c)), None)
        if cmd:
            subprocess.run(
                [cmd, "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
                capture_output=True, text=True, timeout=60
            )
            if os.path.exists(pdf_path):
                return pdf_path
    except Exception:
        pass

    # Método 2: docx2pdf (requiere pip install docx2pdf)
    try:
        from docx2pdf import convert
        convert(docx_path, pdf_path)
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    # Método 3: En Mac, usar el comando nativo
    try:
        subprocess.run(
            ["/Applications/LibreOffice.app/Contents/MacOS/soffice",
             "--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path],
            capture_output=True, text=True, timeout=60
        )
        if os.path.exists(pdf_path):
            return pdf_path
    except Exception:
        pass

    return None


def generar_cotizacion(datos):
    """
    Genera Word y PDF a partir del JSON de datos.
    
    Args:
        datos: dict con estructura definida en SKILL.md paso 6
    
    Returns:
        dict con rutas de archivos y metadata
    """
    # Buscar template
    template_path = buscar_template(datos.get("template_path", ""))
    if not template_path:
        return {"success": False, "error": "Template Word no encontrado"}
    
    project_root = os.path.dirname(os.path.abspath(__file__))  # scripts/
    project_root = os.path.dirname(os.path.dirname(os.path.dirname(os.path.dirname(project_root))))  # raíz del repo
    output_dir = datos.get("output_dir") or os.path.join(project_root, "cotizaciones")
    os.makedirs(output_dir, exist_ok=True)
    
    doc = Document(template_path)
    
    cliente = datos["cliente"]
    evento = datos.get("evento", {})
    lineas = datos["lineas"]
    totales = datos["totales"]
    condiciones = datos.get("condiciones", {})
    numero = datos.get("numero_cotizacion", "___")
    _now = datetime.now()
    fecha = datos.get("fecha", f"{_now.day} de {_now.strftime('%B')} de {_now.year}")
    
    # ── 1. PLACEHOLDERS DEL ENCABEZADO ──
    
    replace_placeholder(doc, "COT-___", f"COT-{numero}")
    replace_placeholder(doc, "DD de mes de AAAA", fecha)
    replace_placeholder(doc, "{nombre_contacto}", cliente.get("nombre_contacto", ""))
    replace_placeholder(doc, "{cargo}", cliente.get("cargo", ""))
    replace_placeholder(doc, "{empresa}", cliente.get("empresa", ""))
    replace_placeholder(doc, "{email}", cliente.get("email", ""))
    replace_placeholder(doc, "{telefono}", cliente.get("telefono", ""))
    replace_placeholder(doc, "{nombre_corto}", cliente.get("nombre_corto", ""))
    replace_placeholder(doc, "{referencia_servicio}", evento.get("referencia", ""))
    replace_placeholder(doc, "{presencial/virtual}", evento.get("tipo", "presencial"))
    replace_placeholder(doc, "{ubicacion_evento}", evento.get("ubicacion", "Bogotá"))
    
    # ── 2. TABLA DE SERVICIOS ──
    # Tabla 0: servicios. Fila 0=título, Fila 1=headers, Filas 2-6=datos
    
    services_table = doc.tables[0]
    
    for i, linea in enumerate(lineas):
        if i >= 5:
            break
        row = services_table.rows[i + 2]
        
        limpiar_texto_celda(row.cells[0], linea.get("numero", i + 1), 
                           size=8, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        limpiar_texto_celda(row.cells[1], linea.get("servicio", ""), 
                           size=7, bold=True)
        limpiar_texto_celda(row.cells[2], linea.get("idioma_origen", ""), size=7)
        limpiar_texto_celda(row.cells[3], linea.get("idioma_destino", ""), size=7)
        limpiar_texto_celda(row.cells[4], linea.get("fechas", ""), size=7)
        limpiar_texto_celda(row.cells[5], linea.get("horario", ""), size=7)
        limpiar_texto_celda(row.cells[6], str(linea.get("cantidad", "")), 
                           size=7, alignment=WD_ALIGN_PARAGRAPH.CENTER)
        limpiar_texto_celda(row.cells[7], linea.get("unidad", ""), size=7)
        limpiar_texto_celda(row.cells[8], formato_cop(linea.get("precio_unitario", 0)), 
                           size=7, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        limpiar_texto_celda(row.cells[9], formato_cop(linea.get("precio_total", 0)), 
                           size=7, bold=True, alignment=WD_ALIGN_PARAGRAPH.RIGHT)
        limpiar_texto_celda(row.cells[10], linea.get("notas_detalle", ""), size=6)
    
    # ── 3. TABLA DE TOTALES ──
    # Los placeholders ya tienen "$ " antes, así que enviamos sin signo
    
    replace_placeholder(doc, "{subtotal}", formato_cop_sin_signo(totales.get("subtotal", 0)))
    
    if totales.get("exento_iva", False):
        replace_placeholder(doc, "{iva}", "EXENTO")
    else:
        replace_placeholder(doc, "{iva}", formato_cop_sin_signo(totales.get("iva", 0)))
    
    replace_placeholder(doc, "{total}", formato_cop_sin_signo(totales.get("total", 0)))
    
    # ── 4. NOTAS Y CONDICIONES ──
    
    notas = condiciones.get("notas", [])
    for i in range(3):
        placeholder = f"{{notas_servicio_{i+1}}}"
        valor = notas[i] if i < len(notas) else ""
        replace_placeholder(doc, placeholder, valor)
    
    replace_placeholder(doc, "{validez_oferta}",
                       condiciones.get("validez_oferta", "Confirmación a más tardar 1 mes antes del servicio"))
    replace_placeholder(doc, "{forma_pago}",
                       condiciones.get("forma_pago", "A 30 días"))
    
    # ── 5. GUARDAR Y CONVERTIR ──
    
    empresa_clean = (cliente.get("empresa", "cliente")
                     .replace(" ", "_").replace("/", "_").replace(".", ""))
    filename = f"COT_{numero}_{empresa_clean}"
    docx_path = os.path.join(output_dir, f"{filename}.docx")
    doc.save(docx_path)
    
    pdf_path = convertir_a_pdf(docx_path, output_dir)
    
    return {
        "success": True,
        "docx_path": docx_path,
        "pdf_path": pdf_path,
        "numero_cotizacion": numero,
        "cliente": cliente.get("empresa", ""),
        "total": totales.get("total", 0),
        "total_formateado": formato_cop(totales.get("total", 0)),
        "num_lineas": len(lineas),
    }


def main():
    parser = argparse.ArgumentParser(description="Generar cotización ML Traductores")
    parser.add_argument("--datos", type=str, help="JSON string con datos")
    parser.add_argument("--archivo", type=str, help="Ruta a archivo JSON")
    args = parser.parse_args()
    
    if args.archivo:
        with open(args.archivo, "r", encoding="utf-8") as f:
            datos = json.load(f)
    elif args.datos:
        datos = json.loads(args.datos)
    else:
        print("ERROR: Debe proveer --datos o --archivo")
        sys.exit(1)
    
    result = generar_cotizacion(datos)
    print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
